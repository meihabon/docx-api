from fastapi import FastAPI, UploadFile, File
from docx import Document
import tempfile

app = FastAPI()


def extract_text(doc):
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


def extract_formatting(doc):
    section = doc.sections[0]

    def pt(x):
        return x.pt if x else None

    # margins (in points)
    margins = {
        "top_pt": pt(section.top_margin),
        "bottom_pt": pt(section.bottom_margin),
        "left_pt": pt(section.left_margin),
        "right_pt": pt(section.right_margin),
    }

    # font detection (best-effort)
    font_name = None
    font_size = None

    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.name:
                font_name = r.font.name
            if r.font.size:
                font_size = r.font.size.pt
            if font_name and font_size:
                break

    # paragraph formatting (first paragraph sample)
    line_spacing = None

    indentation = {
        "first_line_pt": None,
        "left_pt": None,
        "right_pt": None
    }

    if doc.paragraphs:
        p = doc.paragraphs[0]
        fmt = p.paragraph_format

        # line spacing
        if fmt.line_spacing:
            try:
                line_spacing = float(fmt.line_spacing)
            except:
                line_spacing = str(fmt.line_spacing)

        # first line indent
        if fmt.first_line_indent:
            indentation["first_line_pt"] = fmt.first_line_indent.pt

        # left indent
        if fmt.left_indent:
            indentation["left_pt"] = fmt.left_indent.pt

        # right indent
        if fmt.right_indent:
            indentation["right_pt"] = fmt.right_indent.pt

def extract_reference_formatting(doc):
    text = extract_text(doc)
    lines = text.split('\n')
    
    # Find references section
    ref_start = -1
    for i, line in enumerate(lines):
        if 'references' in line.lower() or 'bibliography' in line.lower():
            ref_start = i + 1
            break
    
    has_references_section = ref_start != -1
    
    references = []
    if has_references_section:
        # Collect references from ref_start onwards
        current_ref = []
        for line in lines[ref_start:]:
            line = line.strip()
            if not line:
                if current_ref:
                    references.append(' '.join(current_ref))
                    current_ref = []
            else:
                current_ref.append(line)
        if current_ref:
            references.append(' '.join(current_ref))
    
    reference_count = len(references)
    
    # Check alphabetical order
    is_alphabetically_arranged = True
    if references:
        authors = []
        for ref in references:
            # Extract first author last name
            words = ref.split()
            if words:
                author = words[0].rstrip(',')
                authors.append(author.lower())
        
        sorted_authors = sorted(authors)
        is_alphabetically_arranged = authors == sorted_authors
    
    # Basic APA compliance check (simplified)
    apa_issues = []
    overall_score = 100
    if references:
        for ref in references:
            # Check if starts with author last name
            if not ref or not ref[0].isupper():
                apa_issues.append(f"Reference does not start with author last name: {ref[:50]}...")
                overall_score -= 10
            # Check for year
            if '(' not in ref or ')' not in ref:
                apa_issues.append(f"Reference missing year in parentheses: {ref[:50]}...")
                overall_score -= 10
        if not is_alphabetically_arranged:
            apa_issues.append("References are not in alphabetical order")
            overall_score -= 20
    overall_score = max(0, overall_score)
    
    return {
        "has_references_section": has_references_section,
        "is_alphabetically_arranged": is_alphabetically_arranged,
        "reference_count": reference_count,
        "references": references,
        "apa_7th_compliance": {
            "overall_score": overall_score,
            "issues": apa_issues
        }
    }

@app.post("/convert-docx")
async def convert_docx(file: UploadFile = File(...)):

    if not file.filename.endswith(".docx"):
        return {"error": "Only DOCX allowed"}

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(await file.read())
        path = tmp.name

    doc = Document(path)

    text = extract_text(doc)
    formatting = extract_formatting(doc)
    reference_formatting = extract_reference_formatting(doc)

    return {
        "filename": file.filename,
        "text": text,
        "word_count": len(text.split()),
        "formatting": formatting,
        "reference_formatting": reference_formatting
    }
